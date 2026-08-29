from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from pathlib import Path

OUT = Path('/home/user/Chapitre_4_MIDAS_Benin.docx')
UP = Path('/home/user/uploads')

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(text); r.bold = bold
    if color: r.font.color.rgb = RGBColor(*color)
    for run in p.runs: run.font.name='Times New Roman'; run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman'); run.font.size=Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        c=table.rows[0].cells[i]; set_cell_text(c,h,True,(255,255,255)); shade(c,'059669')
        if widths: c.width=Inches(widths[i])
    for row in rows:
        cells=table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i],str(val))
            if len(table.rows)%2==0: shade(cells[i],'F0FDF4')
    doc.add_paragraph()
    return table

def add_bullet(doc,text,level=0):
    p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2')
    p.add_run(text)
    return p

def add_caption(doc,text):
    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.italic=True; r.font.size=Pt(10)

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); tblHeader=OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)

def set_page_number(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run('Page ')
    fldChar1=OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'),'begin')
    instrText=OxmlElement('w:instrText'); instrText.set(qn('xml:space'),'preserve'); instrText.text='PAGE'
    fldChar2=OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end')
    paragraph._p.append(fldChar1); paragraph._p.append(instrText); paragraph._p.append(fldChar2)

D=Document()
sec=D.sections[0]
sec.top_margin=Inches(.8); sec.bottom_margin=Inches(.75); sec.left_margin=Inches(.85); sec.right_margin=Inches(.85)
styles=D.styles
styles['Normal'].font.name='Times New Roman'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman'); styles['Normal'].font.size=Pt(12)
styles['Normal'].paragraph_format.line_spacing=1.35
styles['Normal'].paragraph_format.space_after=Pt(6)
for name,size,color in [('Heading 1',15,'064E3B'),('Heading 2',13,'047857'),('Heading 3',12,'059669')]:
    st=styles[name]; st.font.name='Times New Roman'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman'); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_before=Pt(14); st.paragraph_format.space_after=Pt(7)
header=sec.header.paragraphs[0]; header.text='MIDAS-Bénin — Chapitre 4 : Mise en œuvre expérimentale et évaluation'; header.alignment=WD_ALIGN_PARAGRAPH.CENTER
for r in header.runs: r.font.name='Times New Roman'; r.font.size=Pt(9); r.font.color.rgb=RGBColor(5,120,87)
set_page_number(sec.footer.paragraphs[0])

p=D.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('CHAPITRE 4 : MISE EN ŒUVRE EXPÉRIMENTALE ET ÉVALUATION DU PROTOTYPE MIDAS-BÉNIN'); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=RGBColor(4,120,87)
D.add_paragraph()

D.add_heading('Introduction', level=1)
D.add_paragraph("Ce chapitre présente la mise en œuvre, l’expérimentation et l’évaluation du prototype MIDAS-Bénin (Mobile IoT Data protection Architecture for e-Services in Benin). L’objectif est de confronter l’architecture à cinq couches proposée au Chapitre 3 à une réalisation concrète, tout en respectant le périmètre académique du travail : absence d’accès institutionnel à l’ANIP, à l’APDP ou aux e-services publics, absence de budget d’infrastructure et emploi exclusif de données synthétiques.")
D.add_paragraph("Le prototype ne prétend donc pas constituer une plateforme nationale opérationnelle. Il représente une preuve de concept à haute fidélité, capable de démontrer les parcours citoyens, la gestion du consentement, la traçabilité, l’enrôlement par NPI simulé, la connexion réelle d’une montre BLE et l’observation de paquets issus de capteurs. Les résultats présentés distinguent systématiquement ce qui a été effectivement implémenté de ce qui reste simulé ou dépendant d’autorisations institutionnelles.")

D.add_heading('4.1 Mise en œuvre du prototype MIDAS-Bénin', level=1)
D.add_heading('4.1.1 Description de la solution implémentée : architecture logique et physique', level=2)
D.add_paragraph("La solution implémentée combine une interface de portefeuille citoyen, une application Android native hôte, un moteur de simulation local et un module BLE réel pour la montre AMIS Watch5GTR. Le choix d’un conteneur Android WebView permet de conserver à l’identique la maquette React initiale tout en donnant accès aux capacités natives du téléphone, notamment la biométrie, les permissions Bluetooth et les services GATT.")
add_table(D,['Couche','Composants réalisés','Statut expérimental'],[
['Couche 1 — Perception','AMIS Watch5GTR réelle, tensiomètre, glucomètre et capteurs agricoles simulés','Montre réelle ; autres appareils simulés'],
['Couche 2 — Réseau / transport','Bluetooth Low Energy, scan Android, GATT, notifications BLE','Réel pour la montre ; LoRaWAN/NB-IoT simulés'],
['Couche 3 — Traitement','Android Keystore, biométrie, consentements, décodage RDFit ciblé, audit local','Partiellement réel ; stockage serveur simulé'],
['Couche 4 — Application','Portefeuille citoyen, identité DID, credentials, e-services, export','Fonctionnel avec données synthétiques'],
['Couche 5 — Gouvernance','Journal hashé, ancrage de démonstration, console de contrôle représentée','Simulation locale ; pas de réseau Iroha institutionnel']
],[1.2,3.6,1.7])
D.add_paragraph("L’architecture physique retenue est volontairement légère. Le téléphone Android joue le rôle de hub de confiance : il héberge l’application, sollicite la biométrie native, réalise le scan BLE, établit la connexion GATT et transmet les événements à l’interface du portefeuille. La montre AMIS est un périphérique de perception qui annonce ses services BLE et émet des paquets propriétaires RDFit. L’interface web embarquée dans l’APK présente les résultats à l’utilisateur.")

D.add_heading('Architecture de déploiement du prototype', level=3)
add_table(D,['Élément','Technologie / rôle','Nature'],[
['Portefeuille citoyen','React, TypeScript, Tailwind CSS, Vite Single File','Interface embarquée dans l’APK'],
['Hôte Android','Kotlin, Android API 31+, WebView','Accès biométrie et BLE'],
['Sécurité locale','AndroidX Biometric, Android Keystore','Authentification locale réelle'],
['Montre IoT','AMIS Watch5GTR, BLE/GATT','Matériel réel'],
['Analyse BLE','nRF Connect, journal HCI BTSnoop','Outils expérimentaux'],
['Registre d’audit','SHA-256 chaîné, métadonnées d’ancrage','Simulation inspirée d’Iroha'],
['Données citoyennes','Données synthétiques et stockage local','Démonstration académique']
],[1.4,3.5,1.6])

D.add_heading('4.1.2 Déroulement de l’implémentation du portefeuille MIDAS-Bénin', level=2)
D.add_paragraph("La réalisation a été organisée en étapes incrémentales afin de valider progressivement les hypothèses du Chapitre 3. Cette démarche a permis de dissocier les fonctions immédiatement réalisables sans infrastructure nationale des fonctions qui nécessiteraient un partenariat avec l’ANIP, l’APDP, un fabricant ou un fournisseur d’e-service.")
add_table(D,['Étape','Mise en œuvre','Résultat obtenu'],[
['1. Réorganisation','Migration en structure React/TypeScript cohérente et compilation Vite','Projet compilable, typage validé'],
['2. Portefeuille','Écrans tableau de bord, identité, consentements, IoT, e-services, audit et paramètres','Parcours citoyen complet'],
['3. Enrôlement','Saisie du NPI de démonstration, résolution locale simulée, DID unique, credential','Un enrôlement par installation'],
['4. Sécurité','Biométrie Android, Keystore, contrôle d’accès local','Déverrouillage natif du portefeuille'],
['5. Audit','Hash SHA-256, historique, export JSON, référence d’ancrage simulée','Traçabilité vérifiable localement'],
['6. BLE réel','Permissions, scan, GATT, abonnement aux notifications de la montre AMIS','Connexion matérielle validée'],
['7. Analyse RDFit','Journal HCI, corrélation commandes/réponses','Identification de trames capteurs ciblées']
],[1.15,3.6,1.75])

D.add_heading('4.1.3 Spécificités de sécurité, de confidentialité et de souveraineté', level=2)
add_bullet(D,"L’application ne collecte ni empreinte digitale ni image faciale. Elle délègue la vérification biométrique au système Android, qui retourne uniquement un résultat de succès ou d’échec.")
add_bullet(D,"Le NPI est utilisé dans un registre local simulé. Le prototype ne contacte pas l’ANIP et ne doit recevoir aucun NPI réel dans le cadre de l’expérimentation.")
add_bullet(D,"L’enrôlement est verrouillé localement après création : un NPI de démonstration donne lieu à un DID unique par installation. Cette règle matérialise l’usage unique dans les limites du prototype local.")
add_bullet(D,"Les commandes RDFit ne sont limitées qu’aux actions observées dans le journal HCI : mesure de rythme cardiaque, mesure SpO2, tension estimée et fonction de recherche/vibration. Aucune commande arbitraire de configuration, de réinitialisation ou de firmware n’est envoyée.")
add_bullet(D,"Les indicateurs santé fournis par la montre sont traités comme des indicateurs de bien-être provenant d’un appareil grand public, et non comme des résultats médicaux certifiés.")

D.add_heading('4.2 Résultats expérimentaux et discussion', level=1)
D.add_heading('4.2.1 Validation fonctionnelle de MIDAS-Bénin : scénarios, captures et commentaires', level=2)
D.add_paragraph("Les tests ont porté sur les fonctions centrales du prototype : enrôlement, contrôle des consentements, gestion d’objets IoT, audit, interface mobile et communication BLE. Les captures intégrées ci-dessous proviennent de l’exécution sur téléphone Android et des outils de diagnostic BLE.")
add_table(D,['Scénario de test','Procédure','Résultat','État'],[
['Enrôlement NPI / DID','Saisie d’un NPI synthétique et acceptation des conditions','Création d’un DID, d’un credential et verrouillage local','Validé en simulation'],
['Consentement','Ouverture d’une demande, accord/refus/révocation','État modifié et événement d’audit créé','Validé'],
['Audit','Recherche, filtres, détail, export JSON','Journal lisible et exportable','Validé'],
['Responsivité','Exécution sur écran Android portrait','Cartes, modales et audit adaptés au mobile','Validé après ajustements CSS'],
['Scan BLE','Recherche de la montre AMIS avec localisation active','Montre détectée avec RSSI d’environ -46 dBm','Validé matériellement'],
['Connexion GATT','Connexion via nRF Connect','Services standards et propriétaires découverts','Validé matériellement'],
['Paquets RDFit','Journal HCI durant les actions santé','Trames de requête et de réponse corrélées','Validé exploratoirement']
],[1.35,3.1,1.75,1.1])

# Screenshots
for img,cap,width in [
    (UP/'Screenshot_20260813_192348.png','Figure 18 — Interface mobile MIDAS-Bénin : panneau de connexion BLE réelle de la montre AMIS.',3.1),
    (UP/'WhatsApp Image 2026-08-13 at 18.44.43.jpeg','Figure 19 — Connexion GATT établie et découverte des services propriétaires de la montre AMIS Watch5GTR.',3.0),
    (UP/'WhatsApp Image 2026-08-07 at 21.24.100.jpeg','Figure 20 — Détection de la montre AMIS Watch5GTR dans nRF Connect avant interrogation des services.',3.0),
]:
    if img.exists():
        p=D.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(img), width=Inches(width))
        add_caption(D,cap)

D.add_heading('Analyse des résultats BLE', level=3)
D.add_paragraph("La connexion GATT directe à la montre AMIS Watch5GTR a été établie avec succès. La montre annonce une connectivité BLE et BR/EDR, une puissance de signal d’environ -46 dBm et des services propriétaires. L’analyse nRF Connect a révélé notamment un canal de type UART propriétaire, composé du service 6e40ab01-b5a3-f393-e0a9-e50e24dcca9e, d’une caractéristique d’écriture 6e40ab02… et d’une caractéristique de notification 6e40ab03…. Des services additionnels AE00, FEE7, 3802 et cc353442-be58-4ea2-876e-11d8d6976366 ont aussi été observés.")
D.add_paragraph("Le journal BTSnoop a ensuite permis de corréler, à l’heure des actions exécutées dans RDFit, les trames de commande et les réponses de la montre. Les résultats ci-dessous doivent être compris comme une observation expérimentale du protocole RDFit propre à la montre testée.")
add_table(D,['Action réalisée','Trame de commande observée','Trame de réponse observée','Valeur déduite'],[
['Fréquence cardiaque','… 0A 20 01 01','… 0A 0A 58','88 bpm'],
['SpO2','… 0A 20 02 01','… 0A 0D 5F','95 %'],
['Tension estimée','… 0A 20 03 01','… 0A 0C 73 47','115/71 mmHg'],
['Recherche de la montre','… 04 01','Accusé de réception observé','Vibration / localisation']
],[1.35,2.15,2.05,1.5])
D.add_paragraph("Ces résultats permettent d’implémenter un premier décodeur RDFit dans l’application Android. Toutefois, ils ne constituent pas une spécification officielle du fabricant : les commandes sont limitées à celles observées, et les valeurs de santé restent des indicateurs de bien-être. Les fonctions de sommeil, pas, batterie, musique, notifications textuelles et mise à jour firmware nécessitent des observations supplémentaires avant toute intégration.")

D.add_heading('4.2.2 Analyse critique du prototype MIDAS-Bénin : apports, limites et perspectives', level=2)
D.add_heading('Apports et avantages', level=3)
add_table(D,['Dimension','Apport observé'],[
['Souveraineté citoyenne','Le citoyen visualise son identité, ses consentements, ses objets et les actions auditables dans une seule interface.'],
['Traçabilité','Les accords, refus, révocations, appairages, exportations et mesures reçues génèrent des événements d’audit.'],
['Compatibilité mobile','L’interface React a été adaptée aux écrans Android ; le conteneur Kotlin donne accès à la biométrie et au BLE.'],
['Interopérabilité progressive','La connexion BLE réelle a été validée malgré l’usage d’un protocole RDFit propriétaire.'],
['Minimisation','Les données sont conditionnées par le consentement et les commandes de la montre sont limitées aux actions observées.'],
['Coût','La preuve de concept repose sur des outils gratuits : Android Studio, Kotlin, React, Vite, nRF Connect et ADB.']
],[1.6,5.4])

D.add_heading('Limites observées', level=3)
add_table(D,['Limite','Conséquence','Mesure de mitigation / perspective'],[
['Absence d’ANIP','Le NPI n’est pas vérifié nationalement','Registre local synthétique ; intégration ANIP conditionnée à une autorisation formelle'],
['Protocole RDFit propriétaire','Toutes les données ne sont pas immédiatement décodables','Analyse HCI ciblée ; n’intégrer que les commandes vérifiées'],
['Montre grand public','SpO2 et tension non certifiées médicalement','Libellé « indicateur de bien-être » ; pas de diagnostic'],
['Pas de backend national','Pas de persistance serveur ni d’Iroha réel','Hash chain locale et simulation d’ancrage ; déploiement futur contrôlé'],
['Unicité locale du DID','L’unicité est limitée à l’installation Android','Registre DID et vérification institutionnelle à prévoir en production'],
['Tests matériels limités','Une seule famille de montre testée','Répéter les essais sur plusieurs modèles BLE et plusieurs téléphones']
],[1.55,2.2,3.25])

D.add_heading('Discussion', level=3)
D.add_paragraph("Les résultats confirment la faisabilité du rôle de hub de confiance attribué au smartphone dans l’architecture MIDAS-Bénin. Le téléphone parvient à relier une interface citoyenne, un mécanisme biométrique Android, une logique de consentement et un objet connecté réel. La connexion à l’AMIS Watch5GTR est particulièrement significative, car elle montre que l’architecture ne repose pas exclusivement sur des données statiques : elle peut interagir avec un périphérique BLE réel et observer ses échanges propriétaires.")
D.add_paragraph("La difficulté principale concerne l’interopérabilité. L’absence de service standard de fréquence cardiaque impose une phase d’observation du protocole RDFit. Cette limite confirme la pertinence de la couche de traitement proposée au Chapitre 3 : elle doit inclure des adaptateurs par fabricant, une validation stricte des formats de paquets et une séparation entre la collecte technique et l’interprétation métier des mesures.")
D.add_paragraph("Enfin, la solution proposée ne doit pas être interprétée comme un dispositif médical. La valeur scientifique du prototype réside dans la démonstration du contrôle citoyen, de la traçabilité, de l’enrôlement et de la connexion sécurisée d’objets IoT, et non dans la validation clinique des capteurs grand public.")

D.add_heading('Conclusion du chapitre', level=1)
D.add_paragraph("Ce chapitre a présenté la mise en œuvre et les résultats expérimentaux de MIDAS-Bénin. Le prototype a permis de valider une interface de portefeuille citoyen, un mécanisme d’enrôlement local fondé sur un NPI de démonstration, la génération d’un DID, la gestion de consentements, l’export et l’audit des actions. L’expérimentation matérielle avec la montre AMIS Watch5GTR a confirmé la détection BLE, la connexion GATT, la découverte de services propriétaires et l’observation de trames RDFit correspondant à des mesures de rythme cardiaque, SpO2, tension estimée et recherche de montre. Les résultats soutiennent la faisabilité conceptuelle de l’architecture à cinq couches, tout en mettant en évidence les prérequis nécessaires à un déploiement institutionnel : API ANIP, gouvernance APDP, infrastructure souveraine, spécifications constructeurs et audits de sécurité.")

D.save(OUT)
print(OUT)
