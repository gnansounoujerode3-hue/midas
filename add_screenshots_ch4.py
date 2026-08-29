from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path

DOC=Path('/home/user/Chapitre_4_MIDAS_Benin.docx')
UP=Path('/home/user/uploads')
doc=Document(DOC)

items=[
('Screenshot_20260815_225928.png','Figure 21 — Tableau de bord citoyen : synthèse de l’état du portefeuille.','Cette vue présente les indicateurs principaux : consentements actifs, nombre total de consentements, objets IoT connectés et activité d’audit. Elle constitue le point d’entrée du citoyen vers les fonctions de contrôle et de supervision.'),
('Screenshot_20260815_225956.png','Figure 22 — Activité récente et accès rapide aux fonctions du portefeuille.','La liste récapitule les événements récents du journal. Les actions rapides permettent d’accéder à l’appairage IoT, à l’identité numérique et aux e-services. Le panneau « Action requise » signale une demande de consentement en attente.'),
('Screenshot_20260815_230005.png','Figure 23 — Alerte de consentement et statut de sécurité.','Cette capture illustre la notification d’une demande de consentement ainsi que le panneau de sécurité. Les mentions de chiffrement, Keystore/StrongBox et audit blockchain sont présentées comme des capacités du prototype ou des simulations documentées.'),
('Screenshot_20260815_230021.png','Figure 24 — Carte d’identité numérique du citoyen.','La carte regroupe le NPI, le DID généré à l’enrôlement, les boutons de copie et les indicateurs de protection. Elle matérialise le portefeuille DID/VC présenté dans l’architecture MIDAS-Bénin.'),
('Screenshot_20260815_231615.png','Figure 25 — Informations personnelles et amorce des credentials.','La page affiche les informations résolues par le registre NPI simulé : naissance, adresse, téléphone, courriel et date d’enrôlement. Les données affichées sont synthétiques dans le cadre expérimental.'),
('Screenshot_20260815_231636.png','Figure 26 — Credential d’identité nationale simulé.','Le credential contient son émetteur de simulation, ses dates de validité et les attributs certifiés. Le bouton « Présenter » simule une présentation sélective à un e-service sans transmettre une identité réelle.'),
('Screenshot_20260815_231656.png','Figure 27 — Synthèse de la gestion des consentements.','Les compteurs distinguent les consentements actifs, les demandes en attente, les décisions historiques et le total. La demande Smart City Cotonou peut être examinée par le citoyen.'),
('Screenshot_20260815_231714.png','Figure 28 — Liste des consentements actifs et révocation.','Chaque consentement affiche le responsable, la finalité, les catégories de données, l’expiration et l’action de révocation. La révocation génère une trace d’audit dans le prototype.'),
('Screenshot_20260815_231731.png','Figure 29 — Gestion des objets IoT et panneau BLE de la montre AMIS.','Cette vue ajoute la connexion BLE réelle de l’AMIS Watch5GTR : recherche, connexion, découverte des canaux GATT et indication que les valeurs provenant de la montre doivent être traitées comme des indicateurs de bien-être.'),
('Screenshot_20260815_231805.png','Figure 30 — Objets médicaux simulés : tensiomètre et glucomètre.','Les cartes présentent le statut, la localisation, la batterie, la dernière mesure et les commandes de synchronisation. Elles démontrent le modèle de gestion d’objets IoT de santé.'),
('Screenshot_20260815_231820.png','Figure 31 — Capteurs agricoles simulés.','La page montre les capteurs de température et d’humidité du sol, leurs dernières mesures, leur batterie et une alerte de batterie faible. Ce scénario illustre l’extension de MIDAS-Bénin au secteur agricole.'),
('Screenshot_20260815_231850.png','Figure 32 — Catalogue des e-services.','La page permet de rechercher et filtrer les e-services. Les indicateurs montrent les services connectés, en attente et disponibles. Les catégories sont adaptées à l’administration, la santé, la finance et l’agriculture.'),
('Screenshot_20260815_231909.png','Figure 33 — E-services connectés et demande en attente.','Les cartes de Dossier Médical Partagé et Smart City Cotonou indiquent les données requises, l’état de connexion et l’action associée. Le statut dépend du consentement simulé du citoyen.'),
('Screenshot_20260815_231922.png','Figure 34 — E-services non connectés.','AgriDigital Bénin et Mobile Money Bénin illustrent une demande d’accès aux données nécessaires. Le bouton « Se connecter » produit une demande de connexion simulée et ne contacte aucun service réel.'),
('Screenshot_20260815_231958.png','Figure 35 — Tableau de bord du journal d’audit.','Cette vue résume le volume d’événements, les ancrages de démonstration, les actions liées aux consentements et aux données. Le bloc Hyperledger Iroha 2 représente une simulation de registre permissionné.'),
('Screenshot_20260815_232019.png','Figure 36 — Recherche et lecture mobile des événements d’audit.','Sur téléphone, le tableau d’audit est rendu sous forme de cartes lisibles. Chaque événement expose action, détails, date et statut d’ancrage. Cette adaptation évite le défilement horizontal sur petit écran.'),
('Screenshot_20260815_232044.png','Figure 37 — Paramètres de sécurité et de notifications.','Le citoyen peut modifier les préférences de notification, choisir une apparence et visualiser l’état des protections. La biométrie est gérée localement par Android ; l’application ne stocke pas de données biométriques brutes.'),
('Screenshot_20260815_232108.png','Figure 38 — Préférences de notification, apparence et droits.','La page met en évidence les interrupteurs de notification, la sélection de thème et les droits de portabilité ou d’effacement. Les options agissent dans le périmètre local du prototype.'),
('Screenshot_20260815_232129.png','Figure 39 — Portabilité, effacement et informations légales.','Les boutons permettent de produire un export JSON et d’enregistrer une demande d’effacement de démonstration. Le bloc légal rappelle le positionnement académique de MIDAS-Bénin et la référence au Livre V du Code du Numérique.')
]

doc.add_page_break()
doc.add_heading('4.2.1.1 Atlas des interfaces expérimentales et description fonctionnelle',level=2)
doc.add_paragraph("Les figures suivantes documentent les écrans effectivement obtenus lors de l’exécution du prototype Android MIDAS-Bénin. Elles constituent des preuves visuelles complémentaires des tests fonctionnels présentés à la section 4.2.1. Chaque description précise la fonction représentée et, lorsque nécessaire, rappelle le périmètre de simulation académique.")

for i,(name,caption,desc) in enumerate(items):
    path=UP/name
    if not path.exists():
        continue
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path),width=Inches(2.35))
    cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=cap.add_run(caption); r.italic=True; r.font.size=Pt(9); r.font.name='Times New Roman'
    d=doc.add_paragraph(); d.paragraph_format.space_after=Pt(10)
    rr=d.add_run('Fonction représentée : '); rr.bold=True
    d.add_run(desc)
    if (i+1)%2==0 and i != len(items)-1:
        doc.add_page_break()

doc.save(DOC)
print(DOC)
