from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT=Path('/home/user/Chapitre_4_MIDAS_Benin_Final.docx')
UP=Path('/home/user/uploads')
D=Document()
sec=D.sections[0]
sec.top_margin=Inches(.8); sec.bottom_margin=Inches(.7); sec.left_margin=Inches(.8); sec.right_margin=Inches(.8)

# Global typography
for name in ['Normal','Heading 1','Heading 2','Heading 3']:
    s=D.styles[name]; s.font.name='Times New Roman'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Times New Roman')
D.styles['Normal'].font.size=Pt(12); D.styles['Normal'].paragraph_format.line_spacing=1.35; D.styles['Normal'].paragraph_format.space_after=Pt(6)
for n,size,col in [('Heading 1',15,'064E3B'),('Heading 2',13,'047857'),('Heading 3',12,'059669')]:
    s=D.styles[n]; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(col); s.paragraph_format.space_before=Pt(14); s.paragraph_format.space_after=Pt(7)

# Header/footer (standard OOXML only for maximum Microsoft Word compatibility)
h=sec.header.paragraphs[0]; h.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=h.add_run('MIDAS-Bénin — Chapitre 4 : Mise en œuvre et résultats expérimentaux'); r.font.name='Times New Roman'; r.font.size=Pt(9); r.font.color.rgb=RGBColor(5,120,87)
f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER
rr=f.add_run('MIDAS-Bénin — Prototype académique'); rr.font.name='Times New Roman'; rr.font.size=Pt(9)

def p(text='', bold_start=None):
    para=D.add_paragraph()
    if bold_start and text.startswith(bold_start):
        para.add_run(bold_start).bold=True; para.add_run(text[len(bold_start):])
    else: para.add_run(text)
    return para

def bullet(text):
    para=D.add_paragraph(style='List Bullet'); para.add_run(text); return para

def shade(cell, color):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),color); tcPr.append(shd)

def table(headers, rows, widths=None):
    t=D.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,hdr in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=hdr; shade(c,'059669'); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for rr in c.paragraphs[0].runs: rr.font.bold=True; rr.font.color.rgb=RGBColor(255,255,255); rr.font.size=Pt(9)
        if widths:c.width=Inches(widths[i])
    for idx,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=str(val); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if idx%2==0: shade(cells[i],'F0FDF4')
            for para in cells[i].paragraphs:
                for rr in para.runs: rr.font.name='Times New Roman'; rr.font.size=Pt(9)
    D.add_paragraph()
    return t

def fig(filename, number, title, comment, width=2.25):
    path=UP/filename
    if not path.exists(): return
    # Les captures sont converties vers JPEG RGB standard pour une compatibilité
    # maximale avec Microsoft Word, y compris les versions anciennes.
    from PIL import Image
    safe_dir=Path('/home/user/docx_safe_images'); safe_dir.mkdir(exist_ok=True)
    safe_path=safe_dir/(path.stem.replace('/', '_') + '.jpg')
    image=Image.open(path).convert('RGB')
    image.save(safe_path, 'JPEG', quality=92, optimize=True)
    par=D.add_paragraph(); par.alignment=WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(str(safe_path),width=Inches(width))
    cap=D.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=cap.add_run(f'Figure {number} — {title}'); run.italic=True; run.font.size=Pt(9)
    desc=D.add_paragraph(); desc.paragraph_format.space_after=Pt(10)
    a=desc.add_run('Commentaire fonctionnel : '); a.bold=True
    desc.add_run(comment)

# Title
par=D.add_paragraph(); par.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=par.add_run('CHAPITRE 4 : MISE EN ŒUVRE, EXPÉRIMENTATION ET ÉVALUATION DE MIDAS-BÉNIN'); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=RGBColor(4,120,87)
D.add_paragraph()

D.add_heading('Introduction',1)
p("Ce chapitre expose la mise en œuvre expérimentale de MIDAS-Bénin (Mobile IoT Data protection Architecture for e-Services in Benin) et l’évaluation du prototype résultant. L’objectif est de confronter l’architecture sécurisée à cinq couches définie au Chapitre 3 à une réalisation concrète : portefeuille citoyen, enrôlement fondé sur un NPI de démonstration, DID, consentements, e-services, objets IoT, audit et connexion Bluetooth Low Energy d’une montre réelle.")
p("Le travail est réalisé dans un cadre académique à ressources limitées. Aucun accès n’a été obtenu auprès de l’ANIP, de l’APDP, de l’ASIN, d’un ministère ou d’un opérateur d’e-service. Les données citoyennes sont synthétiques ; les interfaces institutionnelles et le registre national sont simulés. En revanche, l’interface Android, la biométrie du téléphone, le scan BLE, la connexion GATT à une montre AMIS Watch5GTR et l’analyse des trames RDFit ont été réalisés et testés effectivement.")

D.add_heading('4.1 Mise en œuvre du prototype MIDAS-Bénin',1)
D.add_heading('4.1.1 Présentation de la solution implémentée',2)
p("MIDAS-Bénin est implémenté sous la forme d’une application Android hybride. L’interface citoyenne React/TypeScript est conservée dans l’APK afin de préserver le design fonctionnel initial. Une couche Kotlin native héberge le WebView, gère la biométrie Android, les permissions Bluetooth, la connexion GATT et la communication entre le module BLE et l’interface du portefeuille.")
table(['Composant','Rôle dans le prototype','État'],[
['Interface citoyenne','React, TypeScript, Tailwind CSS, responsive mobile','Implémentée'],
['Conteneur Android','Kotlin, WebView, biométrie, permissions BLE','Implémenté'],
['Enrôlement','NPI de démonstration, résolution locale, DID unique, VC','Implémenté en simulation'],
['Consentements','Accord, refus, révocation, granularité, audit','Implémenté localement'],
['Objets IoT','Cartes IoT, appairage simulé, montre AMIS BLE','Partiellement réel'],
['Audit','Hash chain, export JSON, ancrage de démonstration','Implémenté / simulé'],
['Registre institutionnel','Hyperledger Iroha multi-organisations','Non déployé ; représenté']
],[1.35,3.8,1.3])

D.add_heading('4.1.2 Architecture logique et physique du prototype',2)
p("La couche perception est représentée par des objets IoT simulés et par la montre AMIS Watch5GTR réelle. La couche réseau utilise BLE/GATT pour la montre. Le téléphone constitue le hub de confiance : il sollicite la biométrie, applique les permissions Android, établit la connexion, relaie les événements vers le portefeuille et journalise les actions. La couche application expose les pages citoyennes. La couche de gouvernance est matérialisée par un journal hashé et des références d’ancrage de démonstration.")
table(['Couche IoT','Réalisation MIDAS-Bénin','Niveau de validation'],[
['Perception','Montre AMIS réelle ; tensiomètre, glucomètre et capteurs agricoles de démonstration','Montre validée matériellement'],
['Réseau / transport','BLE, scan Android, GATT, notifications','Validé sur montre AMIS'],
['Traitement','Consentement, DID, audit, WebView↔Kotlin, décodage RDFit ciblé','Partiellement validé'],
['Application','Portefeuille citoyen et e-services simulés','Validé fonctionnellement'],
['Gouvernance','Journal SHA-256 et métadonnées Iroha de démonstration','Validé localement']
],[1.4,3.5,1.55])

D.add_heading('4.1.3 Processus d’enrôlement, DID et identité citoyenne',2)
p("Le parcours d’enrôlement a été revu pour ne demander que le NPI de démonstration. Le nom et le prénom ne sont plus saisis manuellement. Après vérification du format à seize chiffres, un registre local simule la résolution de l’identité associée au NPI. Le système génère ensuite un DID et un credential d’identité. L’enrôlement est conservé localement et verrouillé par installation afin de matérialiser le principe d’usage unique.")
fig('WhatsApp Image 2026-08-15 at 22.39.58.jpeg',1,'Écran initial d’enrôlement MIDAS-Bénin',"Le citoyen saisit uniquement un NPI de démonstration. L’écran indique l’usage unique local, la résolution simulée de l’identité et le DID qui sera généré.")
fig('0.jpeg',2,'Validation locale du NPI et création du DID à usage unique',"Cette vue confirme que l’enrôlement ne sollicite ni nom ni prénom. Le DID est généré une fois, conservé localement et associé au credential d’identité de démonstration.")
fig('Screenshot_20260815_230021.png',3,'Carte DID et NPI du portefeuille citoyen',"La carte d’identité expose le NPI, le DID, les indicateurs de protection et les fonctions de copie. Elle constitue la représentation visuelle du wallet DID/VC.")
fig('Screenshot_20260815_231615.png',4,'Informations personnelles résolues',"Les informations affichées sont résolues par le registre NPI simulé. Elles restent synthétiques dans le prototype, faute d’autorisation d’interroger l’ANIP.")
fig('Screenshot_20260815_231636.png',5,'Credential d’identité nationale simulé',"La carte présente l’émetteur simulé, les dates, les attributs certifiés et la fonction de présentation sélective à un e-service de démonstration.")

D.add_heading('4.1.4 Mise en œuvre de la gestion des consentements',2)
p("Le module de consentement constitue le pivot fonctionnel entre le citoyen, les e-services et les données. Il distingue les consentements actifs, les demandes en attente et les décisions historiques. Le citoyen peut consulter les données demandées, la finalité et le responsable de traitement ; il peut ensuite accorder, refuser ou révoquer un consentement. Chaque transition génère une entrée d’audit locale.")
fig('Screenshot_20260815_231656.png',6,'Tableau de synthèse des consentements',"Les quatre indicateurs montrent les consentements actifs, en attente, révoqués/refusés et le total. La demande Smart City Cotonou illustre une demande à examiner.")
fig('Screenshot_20260815_231714.png',7,'Consentements actifs et révocation',"Chaque consentement présente les données autorisées, l’expiration, le responsable et le bouton de révocation. Cette fonction applique le principe de contrôle permanent du citoyen.")
fig('Screenshot_20260815_230005.png',8,'Action requise et accès rapide au consentement',"Le tableau de bord avertit le citoyen lorsqu’une décision est attendue et fournit un accès direct à la gestion des consentements.")

D.add_heading('4.1.5 Mise en œuvre des objets IoT et de la connexion BLE',2)
p("Le prototype distingue deux niveaux. Le premier est une simulation métier d’objets médicaux et agricoles, destinée à illustrer les parcours IoT. Le second est une expérimentation matérielle réelle avec l’AMIS Watch5GTR. L’application Android demande les permissions Bluetooth et de localisation nécessaires selon le téléphone, détecte la montre, établit la connexion GATT, découvre ses services et souscrit aux caractéristiques de notification identifiées.")
fig('Screenshot_20260815_231731.png',9,'Panneau de connexion BLE réelle de la montre AMIS',"Le panneau propose la recherche, la connexion et la déconnexion de la montre. Il rappelle les canaux GATT surveillés et précise le caractère non médical des indicateurs de bien-être.")
fig('Screenshot_20260815_231805.png',10,'Objets IoT médicaux de démonstration',"Les cartes du tensiomètre et du glucomètre simulent le statut, la batterie, la dernière mesure et la synchronisation. Elles permettent de valider l’expérience utilisateur IoT santé.")
fig('Screenshot_20260815_231820.png',11,'Objets IoT agricoles de démonstration',"Les capteurs de température et d’humidité du sol illustrent l’extension de l’architecture au secteur agricole, y compris l’alerte de batterie faible.")

D.add_heading('4.1.6 Expérimentation de la montre AMIS Watch5GTR',2)
p("La montre AMIS Watch5GTR a été détectée avec un signal BLE voisin de -46 dBm. Une connexion GATT a été établie à l’aide de nRF Connect. Les services observés incluent Generic Access (0x1800), Generic Attribute (0x1801), des services propriétaires AE00, FEE7, 3802, cc353442-be58-4ea2-876e-11d8d6976366 et un canal de type UART propriétaire : 6e40ab01-b5a3-f393-e0a9-e50e24dcca9e avec écriture 6e40ab02… et notifications 6e40ab03…. Cette observation démontre la faisabilité de la couche de transport BLE du prototype.")
fig('WhatsApp Image 2026-08-13 at 18.44.43.jpeg',12,'Services GATT propriétaires de la montre AMIS',"La connexion GATT expose des caractéristiques WRITE, READ et NOTIFY. Les caractéristiques de notification constituent les canaux utilisés pour recevoir des trames de données RDFit.")
table(['Action observée dans RDFit','Trame de commande corrélée','Réponse reçue','Interprétation expérimentale'],[
['Mesure cardiaque','… 0A 20 01 01','… 0A 0A 58','88 bpm'],
['Mesure SpO2','… 0A 20 02 01','… 0A 0D 5F','95 %'],
['Tension estimée','… 0A 20 03 01','… 0A 0C 73 47','115/71 mmHg'],
['Recherche / vibration','… 04 01','Accusé de réception','Commande de localisation de la montre']
],[1.45,1.85,1.8,1.95])
p("L’analyse du journal Bluetooth HCI a permis de corréler les actions exécutées dans RDFit et les paquets GATT. Un décodeur RDFit ciblé a été préparé afin d’identifier ces trames. Les valeurs affichées doivent être considérées comme des indicateurs de bien-être issus d’une montre grand public ; elles ne constituent ni un diagnostic ni une mesure médicale certifiée.")

D.add_heading('4.2 Résultats et discussion',1)
D.add_heading('4.2.1 Résultats des tests de bon fonctionnement',2)
p("Les tests fonctionnels ont été réalisés sur le portefeuille Android, sur les parcours citoyens et sur la connexion à la montre. Le tableau suivant synthétise les résultats obtenus.")
table(['Fonction testée','Procédure','Résultat','Statut'],[
['Enrôlement','Saisie NPI synthétique, validation, DID','DID et VC créés ; identité résolue localement','Validé'],
['Biométrie','Déverrouillage par API AndroidX Biometric','Authentification déléguée au téléphone','Validé'],
['Consentement','Accord, refus, révocation','États mis à jour et audit généré','Validé'],
['E-services','Recherche, filtre, demande d’accès, révocation','Parcours simulés fonctionnels','Validé en simulation'],
['Audit','Recherche, affichage mobile, export','Événements présentés et export JSON','Validé'],
['Export / effacement','Téléchargement et demande d’effacement','Fonctions locales de démonstration','Validé en simulation'],
['Responsivité','Exécution Android portrait','Pages adaptées au petit écran','Validé'],
['BLE montre','Scan, connexion, GATT, notifications','Connexion matérielle établie','Validé'],
['Décodage RDFit','Corrélation HCI de trames santé','Trois formats de réponse identifiés','Exploratoire validé']
],[1.35,2.55,2.1,1.05])

D.add_heading('4.2.2 Présentation commentée des e-services',2)
p("Les e-services sont présentés comme des partenaires de démonstration. Ils ne communiquent avec aucun serveur public réel. Leur rôle est de matérialiser la demande de données, la finalité, la minimisation, le statut de connexion et l’exercice de la révocation.")
fig('Screenshot_20260815_231850.png',13,'Catalogue et filtres des e-services',"La recherche et les filtres par catégorie permettent d’explorer les e-services. Les compteurs indiquent les états connectés, en attente et disponibles.")
fig('Screenshot_20260815_231909.png',14,'E-services connectés et demande en attente',"Le Dossier Médical Partagé et Smart City Cotonou illustrent deux états de relation distincts : connecté et en attente. Les données demandées sont affichées avant toute décision citoyenne.")
fig('Screenshot_20260815_231922.png',15,'E-services non connectés',"AgriDigital Bénin et Mobile Money Bénin présentent une fonction de demande d’accès simulée. Le bouton ne contacte pas les services réels mais reproduit l’état métier attendu.")

D.add_heading('4.2.3 Validation du journal d’audit',2)
p("Le journal d’audit donne au citoyen une visibilité sur les opérations sensibles. Les événements sont représentés par une chaîne de hash SHA-256 dans le prototype ; les références Hyperledger affichées sont des ancrages de démonstration et non des transactions sur un réseau institutionnel réel.")
fig('Screenshot_20260815_231958.png',16,'Vue d’ensemble du journal d’audit',"Le tableau de bord d’audit indique le total des événements, les références d’ancrage de démonstration et les catégories liées aux consentements et aux données.")
fig('Screenshot_20260815_232019.png',17,'Audit adapté à un écran mobile',"Le tableau traditionnel est transformé en cartes sur smartphone afin de préserver la lisibilité des champs Action, Détails, Date et Blockchain.")

D.add_heading('4.2.4 Paramètres, sécurité et exercice des droits',2)
p("La page Paramètres matérialise plusieurs droits citoyens : portabilité, effacement, préférences de notification et apparence. Elle documente également le rôle de la biométrie Android et de la protection locale de clé, tout en précisant le caractère de démonstration de certains indicateurs.")
fig('Screenshot_20260815_232044.png',18,'Sécurité et préférences de notification',"Cette vue présente l’état de sécurité, le contrôle biométrique, la clé de démonstration et les interrupteurs de notification. La biométrie est vérifiée par Android et aucune donnée biométrique brute n’est stockée par MIDAS.")
fig('Screenshot_20260815_232108.png',19,'Apparence et préférences de notification',"Le citoyen peut sélectionner le thème clair, sombre ou système et activer les catégories de notifications. Les préférences sont manipulées dans le périmètre local du prototype.")
fig('Screenshot_20260815_232129.png',20,'Portabilité, effacement et mentions légales',"L’export produit un fichier JSON de démonstration. La suppression enregistre une demande et une trace d’audit ; elle ne doit pas être interprétée comme l’effacement d’un système institutionnel réel.")

D.add_heading('4.2.5 Discussion : apports de MIDAS-Bénin',2)
table(['Axe','Apport observé'],[
['Approche citoyenne','Le portefeuille rassemble identité, consentements, objets IoT, e-services et audit dans une interface unique.'],
['Contrôle des données','Le consentement est visible, modifiable et révocable par le citoyen.'],
['Traçabilité','Les actions significatives génèrent des événements consultables et exportables.'],
['Faisabilité mobile','La biométrie Android et le BLE confirment le rôle du smartphone comme hub de confiance.'],
['Interopérabilité progressive','La connexion GATT à une montre propriétaire démontre qu’un adaptateur par constructeur peut compléter les protocoles standards.'],
['Coût et reproductibilité','Les outils utilisés sont accessibles : Android Studio, Kotlin, React, Vite, nRF Connect et ADB.']
],[1.65,5.25])

D.add_heading('4.2.6 Limites et perspectives',2)
table(['Limite','Impact','Perspective'],[
['ANIP non accessible','Pas de vérification réelle du NPI','API sécurisée et accord institutionnel'],
['APDP et Iroha non déployés','Audit non opposable juridiquement','Réseau permissionné multi-organisations'],
['RDFit propriétaire','Décodage limité aux trames observées','Spécification fournisseur ou analyse contrôlée élargie'],
['Montre grand public','Données santé non cliniques','Utilisation informative ; pas de diagnostic'],
['Données synthétiques','Pas de validation sur population réelle','Pilote avec consentement, analyse d’impact et gouvernance'],
['Stockage local','Persistance limitée à l’appareil','Backend souverain, chiffrement au repos, HSM et rotation de clés']
],[1.55,2.25,3.1])
p("Les résultats confirment la pertinence de l’architecture proposée au Chapitre 3, mais ils ne constituent pas une validation d’un déploiement national. Une mise en production nécessiterait notamment une analyse d’impact sur la protection des données, une validation juridique, une gouvernance multi-acteurs, un hébergement souverain et des audits de sécurité indépendants.")

D.add_heading('Conclusion du chapitre',1)
p("Ce chapitre a présenté la mise en œuvre et l’évaluation expérimentale de MIDAS-Bénin. Le prototype valide les parcours fondamentaux : enrôlement local via NPI de démonstration, DID, credential, consentement, révocation, audit, e-services simulés, export et interface responsive. Il valide également une connexion BLE/GATT réelle avec la montre AMIS Watch5GTR et l’observation de trames RDFit correspondant à des actions de mesure et de pilotage ciblées. Ces résultats soutiennent la faisabilité d’un portefeuille citoyen IoT centré sur le consentement et la traçabilité, tout en soulignant les limites institutionnelles, cryptographiques et médicales qui devront être traitées avant tout déploiement opérationnel.")

D.save(OUT)
print(OUT)
